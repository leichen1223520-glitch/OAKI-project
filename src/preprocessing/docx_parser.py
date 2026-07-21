"""
docx_parser.py — 解析41座设施个厂资料 (.docx)，提取设备能耗特征数据

输出字段（PlantEquipmentData）：
  plant_name         厂站名称（从文件名清洗）
  source_file        原始文件路径
  intake_pump_power_kw    进水泵单机功率 (kW)
  intake_pump_count       进水泵运行台数
  blower_power_kw         鼓风机单机功率 (kW)
  blower_count            鼓风机运行台数
  total_blower_power_kw   鼓风机总装机功率 (kW)
  dewater_type            脱水机类型（板框/离心/带式/叠螺）
  dewater_power_kw        脱水机单机功率 (kW)
  dewater_count           脱水机台数
  do_aerobic_min          好氧区DO下限 (mg/L)
  do_aerobic_max          好氧区DO上限 (mg/L)
  sludge_moisture_out     出泥含水率 (%)
  has_drying              是否有干化设备 (bool)
  has_mbr                 是否含MBR工艺 (bool)
  notes                   备注/解析警告
"""
from __future__ import annotations

import re
import glob
import json
from dataclasses import dataclass, field, asdict
from pathlib import Path
from typing import Optional, List

try:
    from docx import Document
except ImportError:
    raise ImportError("请安装 python-docx: pip install python-docx")


# ─── 数据容器 ───────────────────────────────────────────────────────────────

@dataclass
class PlantEquipmentData:
    plant_name: str = ""
    source_file: str = ""
    intake_pump_power_kw: Optional[float] = None
    intake_pump_count: Optional[int] = None
    blower_power_kw: Optional[float] = None
    blower_count: Optional[int] = None
    total_blower_power_kw: Optional[float] = None
    dewater_type: Optional[str] = None
    dewater_power_kw: Optional[float] = None
    dewater_count: Optional[int] = None
    do_aerobic_min: Optional[float] = None
    do_aerobic_max: Optional[float] = None
    sludge_moisture_out: Optional[float] = None
    has_drying: bool = False
    has_mbr: bool = False
    notes: List[str] = field(default_factory=list)


# ─── 辅助正则 ────────────────────────────────────────────────────────────────

_RE_POWER = re.compile(r'(\d+\.?\d*)\s*kW', re.IGNORECASE)
_RE_COUNT = re.compile(r'(\d+)\s*台')
_RE_COUNT_RUN = re.compile(r'(\d+)\s*用\s*\d+\s*备|运行台数[：:]\s*(\d+)')
_RE_DO_RANGE = re.compile(r'(\d+\.?\d*)\s*[~～\-–]\s*(\d+\.?\d*)\s*(?:mg/L)?')
_RE_MOISTURE = re.compile(r'(\d+\.?\d*)\s*[%％]')
_RE_FLOAT = re.compile(r'(\d+\.?\d*)')


def _extract_first_float(text: str) -> Optional[float]:
    m = _RE_FLOAT.search(text)
    return float(m.group(1)) if m else None


def _extract_power(text: str) -> Optional[float]:
    """提取文本中第一个功率值 (kW)"""
    m = _RE_POWER.search(text)
    return float(m.group(1)) if m else None


def _extract_count(text: str) -> Optional[int]:
    """提取运行台数"""
    # 优先匹配 "X用Y备"
    for pat in [
        re.compile(r'(\d+)\s*[用台]\s*\d+\s*备'),
        re.compile(r'运行台数[：:]\s*(\d+)'),
        re.compile(r'^(\d+)$'),
        re.compile(r'共\s*(\d+)\s*台'),
        re.compile(r'(\d+)\s*台'),
    ]:
        m = pat.search(text)
        if m:
            return int(m.group(1))
    return None


def _extract_do(text: str):
    """提取 DO 范围 (min, max)"""
    m = _RE_DO_RANGE.search(text)
    if m:
        return float(m.group(1)), float(m.group(2))
    single = re.search(r'DO[：:\s]*(\d+\.?\d*)', text)
    if single:
        v = float(single.group(1))
        return v, v
    return None, None


def _extract_moisture(text: str) -> Optional[float]:
    """提取含水率 %（取第一个合理值）"""
    for m in _RE_MOISTURE.finditer(text):
        v = float(m.group(1))
        if 50 <= v <= 99:
            return v
    return None


# ─── 文件名 → 厂站名 ─────────────────────────────────────────────────────────

def _clean_plant_name(filepath: str) -> str:
    name = Path(filepath).stem
    # 去除常见前缀
    prefixes = [
        r'^附件：\d+\.?排水处咨询工作所需资料[(-（]?',
        r'^排水处咨询工作所需资料[(-（]?',
        r'^.*?资料[(-（]',
    ]
    for pat in prefixes:
        name = re.sub(pat, '', name).strip('） )(（-')
        if name:
            break
    # 特殊清理
    name = re.sub(r'-排水处咨询工作所需资料$', '', name).strip()
    name = re.sub(r'排水处咨询工作所需资料$', '', name).strip()
    name = re.sub(r'\(\d+\)$', '', name).strip()  # 去除 (1) (2)
    # 如果名称中含有"厂"字，提取到厂为止
    if not name:
        name = Path(filepath).stem
    return name


# ─── 主解析函数 ───────────────────────────────────────────────────────────────

def parse_plant_docx(filepath: str) -> PlantEquipmentData:
    """解析单个厂站资料 docx，返回 PlantEquipmentData"""
    result = PlantEquipmentData(
        plant_name=_clean_plant_name(filepath),
        source_file=str(filepath),
    )

    try:
        doc = Document(filepath)
    except Exception as e:
        result.notes.append(f"文件读取失败: {e}")
        return result

    # 将所有段落文本和表格单元格文本合并，按节分区处理
    sections = {
        'intake': [],
        'blower': [],
        'aeration': [],
        'dewater': [],
        'drying': [],
        'mbr': [],
    }

    current_section = 'other'
    section_map = {
        '进水': 'intake', '提升': 'intake',
        '鼓风机': 'blower', '风机': 'blower',
        '曝气': 'aeration',
        '脱水': 'dewater', '浓缩': 'dewater', '污泥': 'dewater',
        '干化': 'drying',
        'MBR': 'mbr', 'mbr': 'mbr',
    }

    all_blocks: List[str] = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            all_blocks.append(t)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t and len(t) > 2:
                    all_blocks.append(t)

    for block in all_blocks:
        # 更新当前节
        for kw, sec in section_map.items():
            if kw in block and len(block) < 30:
                current_section = sec
                break

        for sec_kw, sec in sections.items():
            pass  # just collect

        if current_section in sections:
            sections[current_section].append(block)

        # MBR 检测
        if re.search(r'MBR|膜生物反应器|超滤膜|UF膜', block, re.IGNORECASE):
            result.has_mbr = True

        # 干化设备
        if '干化' in block and '无' not in block[:5]:
            if any(kw in block for kw in ['kW', '功率', '蒸发', '桨叶', '低温']):
                result.has_drying = True

    # ── 进水泵 ──────────────────────────────────────────────────────────────
    for block in sections['intake']:
        if '进水泵' in block or '提升泵' in block:
            if result.intake_pump_power_kw is None:
                result.intake_pump_power_kw = _extract_power(block)
            if result.intake_pump_count is None:
                result.intake_pump_count = _extract_count(block)

    # 也从表格里找（Table 0 第一行通常是进水泵）
    try:
        t0 = doc.tables[0]
        for row in t0.rows:
            row_text = ' '.join(c.text.strip() for c in row.cells)
            if '进水泵' in row_text:
                if result.intake_pump_power_kw is None:
                    result.intake_pump_power_kw = _extract_power(row_text)
                if result.intake_pump_count is None:
                    result.intake_pump_count = _extract_count(row_text)
                break
    except Exception:
        pass

    # ── 鼓风机 ──────────────────────────────────────────────────────────────
    for block in sections['blower']:
        if result.blower_power_kw is None:
            result.blower_power_kw = _extract_power(block)
        if result.blower_count is None:
            result.blower_count = _extract_count(block)

    # 也从 Table 1 读取（通常是风机表）
    try:
        t1 = doc.tables[1]
        for row in t1.rows:
            row_text = ' '.join(c.text.strip() for c in row.cells)
            if any(kw in row_text for kw in ['鼓风机', '离心风机', '风机']):
                if result.blower_power_kw is None:
                    result.blower_power_kw = _extract_power(row_text)
                if result.blower_count is None:
                    result.blower_count = _extract_count(row_text)
                break
    except Exception:
        pass

    # ── 计算总鼓风机装机功率 ─────────────────────────────────────────────────
    if result.blower_power_kw and result.blower_count:
        result.total_blower_power_kw = result.blower_power_kw * result.blower_count
    elif result.blower_power_kw:
        result.total_blower_power_kw = result.blower_power_kw  # 单台

    # ── 曝气/DO ─────────────────────────────────────────────────────────────
    for block in all_blocks:
        if ('好氧' in block or '曝气' in block) and 'DO' in block:
            lo, hi = _extract_do(block)
            if lo is not None and result.do_aerobic_min is None:
                # 过滤不合理值
                if 0.1 <= lo <= 6.0 and lo <= hi <= 8.0:
                    result.do_aerobic_min = lo
                    result.do_aerobic_max = hi

    # ── 脱水机 ──────────────────────────────────────────────────────────────
    dewater_keywords = {
        '板框': '板框压滤机',
        '带式': '带式压滤机',
        '离心': '离心式脱水机',
        '叠螺': '叠螺脱水机',
    }
    for block in all_blocks:
        if '脱水' in block or '离心' in block:
            if result.dewater_type is None:
                for kw, dtype in dewater_keywords.items():
                    if kw in block:
                        result.dewater_type = dtype
                        break
            if result.dewater_power_kw is None:
                p = _extract_power(block)
                if p and p > 1:  # 过滤过小的值
                    result.dewater_power_kw = p
            if result.dewater_count is None:
                c = _extract_count(block)
                if c and 1 <= c <= 20:
                    result.dewater_count = c

    # ── 含水率 ──────────────────────────────────────────────────────────────
    for block in all_blocks:
        if '出泥含水率' in block or '含水率' in block:
            m = _extract_moisture(block)
            if m and result.sludge_moisture_out is None:
                result.sludge_moisture_out = m

    return result


# ─── 批量解析 ────────────────────────────────────────────────────────────────

def parse_all_plants(data_dir: str = "data/raw/41座设施各厂资料") -> List[PlantEquipmentData]:
    """解析目录中所有 .docx 文件"""
    files = sorted(glob.glob(f"{data_dir}/*.docx"))
    results = []
    for f in files:
        rec = parse_plant_docx(f)
        results.append(rec)
    return results


def to_records(plants: List[PlantEquipmentData]) -> List[dict]:
    """转为 dict 列表（便于 pandas/JSON 导出）"""
    return [asdict(p) for p in plants]


# ─── CLI 入口 ────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    import sys, csv, io
    data_dir = sys.argv[1] if len(sys.argv) > 1 else "data/raw/41座设施各厂资料"
    plants = parse_all_plants(data_dir)
    records = to_records(plants)

    # 输出 JSON
    out_json = "data/processed/facilities_equipment.json"
    with open(out_json, "w", encoding="utf-8") as f:
        json.dump(records, f, ensure_ascii=False, indent=2)
    print(f"已写入 {out_json}（{len(records)} 条记录）")

    # 输出 CSV（排除 notes 字段）
    out_csv = "data/processed/facilities_equipment.csv"
    flat_keys = [k for k in records[0].keys() if k != "notes"]
    with open(out_csv, "w", newline="", encoding="utf-8-sig") as f:
        writer = csv.DictWriter(f, fieldnames=flat_keys, extrasaction="ignore")
        writer.writeheader()
        writer.writerows(records)
    print(f"已写入 {out_csv}（{len(records)} 条记录）")

    # 简要汇总
    parsed_blower = sum(1 for r in plants if r.blower_power_kw)
    parsed_do = sum(1 for r in plants if r.do_aerobic_min)
    print(f"\n解析摘要：")
    print(f"  总文件数：{len(plants)}")
    print(f"  成功提取鼓风机功率：{parsed_blower}/{len(plants)}")
    print(f"  成功提取好氧DO：{parsed_do}/{len(plants)}")
    print(f"  含MBR工艺：{sum(1 for r in plants if r.has_mbr)}")
    print(f"  含干化设备：{sum(1 for r in plants if r.has_drying)}")
